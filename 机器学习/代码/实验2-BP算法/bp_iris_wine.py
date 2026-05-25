from __future__ import annotations

import csv
import math
from dataclasses import dataclass
from pathlib import Path

import numpy as np
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent.parent.parent
DATA_DIR = ROOT / "数据集" / "实验二数据集" / "数据集"
OUT_DIR = ROOT / "输出结果" / "实验2-bp_outputs"
OUT_DIR.mkdir(parents=True, exist_ok=True)
REPORT_DIR = ROOT / "实验报告" / "实验二"
REPORT_OUT = REPORT_DIR / "实验2-BP算法实践-已完成（鸢尾花+葡萄酒）.docx"


@dataclass
class RunResult:
    name: str
    train_acc: float
    test_acc: float
    final_loss: float
    epochs: int
    hidden_size: int
    learning_rate: float
    train_size: int
    test_size: int
    confusion: np.ndarray
    labels: list[str]
    loss_history: list[float]
    acc_history: list[float]


class BPClassifier:
    def __init__(self, n_features: int, n_hidden: int, n_outputs: int, learning_rate: float, seed: int = 42):
        rng = np.random.default_rng(seed)
        self.w1 = rng.normal(0, math.sqrt(2 / (n_features + n_hidden)), size=(n_features, n_hidden))
        self.b1 = np.zeros((1, n_hidden))
        self.w2 = rng.normal(0, math.sqrt(2 / (n_hidden + n_outputs)), size=(n_hidden, n_outputs))
        self.b2 = np.zeros((1, n_outputs))
        self.lr = learning_rate

    @staticmethod
    def sigmoid(z: np.ndarray) -> np.ndarray:
        return 1 / (1 + np.exp(-np.clip(z, -40, 40)))

    @staticmethod
    def softmax(z: np.ndarray) -> np.ndarray:
        z = z - z.max(axis=1, keepdims=True)
        e = np.exp(z)
        return e / e.sum(axis=1, keepdims=True)

    def forward(self, x: np.ndarray) -> tuple[np.ndarray, np.ndarray]:
        h = self.sigmoid(x @ self.w1 + self.b1)
        y = self.softmax(h @ self.w2 + self.b2)
        return h, y

    def fit(self, x: np.ndarray, y: np.ndarray, epochs: int) -> tuple[list[float], list[float]]:
        losses: list[float] = []
        accs: list[float] = []
        n = x.shape[0]
        for _ in range(epochs):
            h, y_hat = self.forward(x)
            loss = -np.sum(y * np.log(y_hat + 1e-12)) / n
            dz2 = (y_hat - y) / n
            dw2 = h.T @ dz2
            db2 = dz2.sum(axis=0, keepdims=True)
            dh = dz2 @ self.w2.T
            dz1 = dh * h * (1 - h)
            dw1 = x.T @ dz1
            db1 = dz1.sum(axis=0, keepdims=True)
            self.w2 -= self.lr * dw2
            self.b2 -= self.lr * db2
            self.w1 -= self.lr * dw1
            self.b1 -= self.lr * db1
            pred = y_hat.argmax(axis=1)
            true = y.argmax(axis=1)
            losses.append(float(loss))
            accs.append(float((pred == true).mean()))
        return losses, accs

    def predict(self, x: np.ndarray) -> np.ndarray:
        return self.forward(x)[1].argmax(axis=1)


def read_iris() -> tuple[np.ndarray, np.ndarray, list[str]]:
    path = DATA_DIR / "iris.csv"
    rows = list(csv.DictReader(path.open("r", encoding="utf-8-sig")))
    labels = sorted({row["Species"] for row in rows})
    label_to_id = {name: i for i, name in enumerate(labels)}
    x = np.array([[float(row[c]) for c in ["Sepal.Length", "Sepal.Width", "Petal.Length", "Petal.Width"]] for row in rows])
    y = np.array([label_to_id[row["Species"]] for row in rows], dtype=int)
    return x, y, labels


def read_wine() -> tuple[np.ndarray, np.ndarray, list[str]]:
    """Read red wine quality dataset, group quality into 3 classes."""
    path = DATA_DIR / "wine+quality" / "winequality-red.csv"
    rows = list(csv.DictReader(path.open("r", encoding="utf-8-sig"), delimiter=";"))
    feature_cols = [
        "fixed acidity", "volatile acidity", "citric acid", "residual sugar",
        "chlorides", "free sulfur dioxide", "total sulfur dioxide", "density",
        "pH", "sulphates", "alcohol"
    ]
    x = np.array([[float(row[c]) for c in feature_cols] for row in rows])
    quality = np.array([int(row["quality"]) for row in rows], dtype=int)
    # Group quality: 3-4 -> low(0), 5-6 -> medium(1), 7-8 -> high(2)
    y = np.zeros(len(quality), dtype=int)
    y[(quality >= 5) & (quality <= 6)] = 1
    y[quality >= 7] = 2
    labels = ["低品质(3-4)", "中等品质(5-6)", "高品质(7-8)"]
    return x, y, labels


def stratified_split(y: np.ndarray, test_ratio: float, seed: int) -> tuple[np.ndarray, np.ndarray]:
    rng = np.random.default_rng(seed)
    train_idx: list[int] = []
    test_idx: list[int] = []
    for cls in np.unique(y):
        idx = np.where(y == cls)[0]
        rng.shuffle(idx)
        n_test = max(1, int(round(len(idx) * test_ratio)))
        test_idx.extend(idx[:n_test].tolist())
        train_idx.extend(idx[n_test:].tolist())
    return np.array(train_idx), np.array(test_idx)


def standardize(train_x: np.ndarray, test_x: np.ndarray) -> tuple[np.ndarray, np.ndarray]:
    mean = train_x.mean(axis=0, keepdims=True)
    std = train_x.std(axis=0, keepdims=True)
    std[std == 0] = 1
    return (train_x - mean) / std, (test_x - mean) / std


def one_hot(y: np.ndarray, classes: int) -> np.ndarray:
    out = np.zeros((len(y), classes))
    out[np.arange(len(y)), y] = 1
    return out


def confusion_matrix(true: np.ndarray, pred: np.ndarray, classes: int) -> np.ndarray:
    matrix = np.zeros((classes, classes), dtype=int)
    for t, p in zip(true, pred):
        matrix[t, p] += 1
    return matrix


def run_dataset(name: str, x: np.ndarray, y: np.ndarray, labels: list[str],
                hidden: int, lr: float, epochs: int, test_ratio: float, seed: int) -> RunResult:
    train_idx, test_idx = stratified_split(y, test_ratio, seed)
    train_x, test_x = standardize(x[train_idx], x[test_idx])
    train_y, test_y = y[train_idx], y[test_idx]
    model = BPClassifier(train_x.shape[1], hidden, len(labels), lr, seed)
    loss_history, acc_history = model.fit(train_x, one_hot(train_y, len(labels)), epochs)
    train_pred = model.predict(train_x)
    test_pred = model.predict(test_x)
    train_acc = float((train_pred == train_y).mean())
    test_acc = float((test_pred == test_y).mean())
    cm = confusion_matrix(test_y, test_pred, len(labels))
    return RunResult(
        name=name,
        train_acc=train_acc,
        test_acc=test_acc,
        final_loss=loss_history[-1],
        epochs=epochs,
        hidden_size=hidden,
        learning_rate=lr,
        train_size=len(train_idx),
        test_size=len(test_idx),
        confusion=cm,
        labels=labels,
        loss_history=loss_history,
        acc_history=acc_history,
    )


def get_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def draw_training_curve(result: RunResult, path: Path) -> None:
    width, height = 900, 520
    margin = 70
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font = get_font(22)
    small = get_font(16)
    draw.text((width // 2 - 140, 15), f"{result.name} BP训练曲线", fill=(20, 20, 20), font=font)
    x0, y0 = margin, height - margin
    x1, y1 = width - margin, margin
    draw.line((x0, y0, x1, y0), fill=(0, 0, 0), width=2)
    draw.line((x0, y0, x0, y1), fill=(0, 0, 0), width=2)
    draw.text((x1 - 40, y0 + 14), "epoch", fill=(0, 0, 0), font=small)
    draw.text((9, y1 - 10), "loss", fill=(0, 0, 0), font=small)
    values = np.array(result.loss_history)
    max_v, min_v = float(values.max()), float(values.min())
    span = max(max_v - min_v, 1e-9)
    sample = np.linspace(0, len(values) - 1, min(300, len(values))).astype(int)
    points = []
    for i in sample:
        px = x0 + (x1 - x0) * i / (len(values) - 1)
        py = y0 - (y0 - y1) * (values[i] - min_v) / span
        points.append((px, py))
    draw.line(points, fill=(35, 102, 170), width=3)
    for ratio in np.linspace(0, 1, 5):
        ytick = y0 - (y0 - y1) * ratio
        val = min_v + span * ratio
        draw.line((x0 - 5, ytick, x0, ytick), fill=(0, 0, 0), width=1)
        draw.text((6, ytick - 10), f"{val:.3f}", fill=(0, 0, 0), font=small)
    img.save(path)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    set_run_font(run, "宋体", 12)


def set_run_font(run, font_name: str, size_pt: float) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)


def add_paragraph(doc: Document, text: str, first_line: bool = True) -> None:
    p = doc.add_paragraph()
    if first_line:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, "宋体", 12)


def add_heading_text(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    set_run_font(run, "宋体", 14 if level == 1 else 12)


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "000000")
        borders.append(tag)
    tbl_pr.append(borders)

def main() -> None:
    # Load datasets
    iris_x, iris_y, iris_labels = read_iris()
    wine_x, wine_y, wine_labels = read_wine()

    print(f"Iris: {iris_x.shape}, classes={iris_labels}, distribution={np.bincount(iris_y)}")
    print(f"Wine: {wine_x.shape}, classes={wine_labels}, distribution={np.bincount(wine_y)}")

    # Run experiments
    iris_result = run_dataset(
        "鸢尾花数据集", iris_x, iris_y, iris_labels,
        hidden=10, lr=0.08, epochs=3500, test_ratio=0.30, seed=7,
    )
    wine_result = run_dataset(
        "葡萄酒品质数据集", wine_x, wine_y, wine_labels,
        hidden=16, lr=0.05, epochs=5000, test_ratio=0.30, seed=7,
    )
    results = [iris_result, wine_result]

    # Draw loss curves
    curve_paths = []
    for result in results:
        path = OUT_DIR / f"{result.name}_loss.png"
        draw_training_curve(result, path)
        curve_paths.append(path)

    # Build and save report
    build_report(results, curve_paths, REPORT_OUT)

    # Print results
    for result in results:
        print(f"\n{'='*60}")
        print(f"{result.name}:")
        print(f"  train_acc={result.train_acc:.4f}, test_acc={result.test_acc:.4f}, loss={result.final_loss:.4f}")
        print(f"  confusion matrix:")
        print(result.confusion)
    print(f"\nReport saved to: {REPORT_OUT}")


if __name__ == "__main__":
    main()
