from __future__ import annotations

import csv
import math
from dataclasses import dataclass
from pathlib import Path

import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
DATA_DIR = ROOT / "_dataset_zip_preview" / "数据集"
OUT_DIR = ROOT / "bp_outputs"
OUT_DIR.mkdir(exist_ok=True)


@dataclass
class RunResult:
    name: str
    train_acc: float
    test_acc: float
    final_loss: float
    epochs: int
    hidden_size: int
    learning_rate: float
    train_size: int
    test_size: int
    confusion: np.ndarray
    labels: list[str]
    loss_history: list[float]
    acc_history: list[float]


class BPClassifier:
    def __init__(self, n_features: int, n_hidden: int, n_outputs: int, learning_rate: float, seed: int = 42):
        rng = np.random.default_rng(seed)
        self.w1 = rng.normal(0, math.sqrt(2 / (n_features + n_hidden)), size=(n_features, n_hidden))
        self.b1 = np.zeros((1, n_hidden))
        self.w2 = rng.normal(0, math.sqrt(2 / (n_hidden + n_outputs)), size=(n_hidden, n_outputs))
        self.b2 = np.zeros((1, n_outputs))
        self.lr = learning_rate

    @staticmethod
    def sigmoid(z: np.ndarray) -> np.ndarray:
        return 1 / (1 + np.exp(-np.clip(z, -40, 40)))

    @staticmethod
    def softmax(z: np.ndarray) -> np.ndarray:
        z = z - z.max(axis=1, keepdims=True)
        e = np.exp(z)
        return e / e.sum(axis=1, keepdims=True)

    def forward(self, x: np.ndarray) -> tuple[np.ndarray, np.ndarray]:
        h = self.sigmoid(x @ self.w1 + self.b1)
        y = self.softmax(h @ self.w2 + self.b2)
        return h, y

    def fit(self, x: np.ndarray, y: np.ndarray, epochs: int) -> tuple[list[float], list[float]]:
        losses: list[float] = []
        accs: list[float] = []
        n = x.shape[0]
        for _ in range(epochs):
            h, y_hat = self.forward(x)
            loss = -np.sum(y * np.log(y_hat + 1e-12)) / n
            dz2 = (y_hat - y) / n
            dw2 = h.T @ dz2
            db2 = dz2.sum(axis=0, keepdims=True)
            dh = dz2 @ self.w2.T
            dz1 = dh * h * (1 - h)
            dw1 = x.T @ dz1
            db1 = dz1.sum(axis=0, keepdims=True)
            self.w2 -= self.lr * dw2
            self.b2 -= self.lr * db2
            self.w1 -= self.lr * dw1
            self.b1 -= self.lr * db1
            pred = y_hat.argmax(axis=1)
            true = y.argmax(axis=1)
            losses.append(float(loss))
            accs.append(float((pred == true).mean()))
        return losses, accs

    def predict(self, x: np.ndarray) -> np.ndarray:
        return self.forward(x)[1].argmax(axis=1)


def read_iris() -> tuple[np.ndarray, np.ndarray, list[str]]:
    path = DATA_DIR / "iris.csv"
    rows = list(csv.DictReader(path.open("r", encoding="utf-8-sig")))
    labels = sorted({row["Species"] for row in rows})
    label_to_id = {name: i for i, name in enumerate(labels)}
    x = np.array([[float(row[c]) for c in ["Sepal.Length", "Sepal.Width", "Petal.Length", "Petal.Width"]] for row in rows])
    y = np.array([label_to_id[row["Species"]] for row in rows], dtype=int)
    return x, y, labels


def read_watermelon() -> tuple[np.ndarray, np.ndarray, list[str]]:
    path = DATA_DIR / "3.3.csv"
    rows: list[list[str]] = []
    with path.open("r", encoding="utf-8-sig") as f:
        for line in f:
            if line.strip():
                rows.append(line.strip().split(","))
    x = np.array([[float(r[1]), float(r[2])] for r in rows])
    y = np.array([int(r[3]) for r in rows], dtype=int)
    return x, y, ["bad", "good"]


def stratified_split(y: np.ndarray, test_ratio: float, seed: int) -> tuple[np.ndarray, np.ndarray]:
    rng = np.random.default_rng(seed)
    train_idx: list[int] = []
    test_idx: list[int] = []
    for cls in np.unique(y):
        idx = np.where(y == cls)[0]
        rng.shuffle(idx)
        n_test = max(1, int(round(len(idx) * test_ratio)))
        test_idx.extend(idx[:n_test].tolist())
        train_idx.extend(idx[n_test:].tolist())
    return np.array(train_idx), np.array(test_idx)


def standardize(train_x: np.ndarray, test_x: np.ndarray) -> tuple[np.ndarray, np.ndarray]:
    mean = train_x.mean(axis=0, keepdims=True)
    std = train_x.std(axis=0, keepdims=True)
    std[std == 0] = 1
    return (train_x - mean) / std, (test_x - mean) / std


def one_hot(y: np.ndarray, classes: int) -> np.ndarray:
    out = np.zeros((len(y), classes))
    out[np.arange(len(y)), y] = 1
    return out


def confusion_matrix(true: np.ndarray, pred: np.ndarray, classes: int) -> np.ndarray:
    matrix = np.zeros((classes, classes), dtype=int)
    for t, p in zip(true, pred):
        matrix[t, p] += 1
    return matrix


def run_dataset(name: str, x: np.ndarray, y: np.ndarray, labels: list[str], hidden: int, lr: float, epochs: int, test_ratio: float, seed: int) -> RunResult:
    train_idx, test_idx = stratified_split(y, test_ratio, seed)
    train_x, test_x = standardize(x[train_idx], x[test_idx])
    train_y, test_y = y[train_idx], y[test_idx]
    model = BPClassifier(train_x.shape[1], hidden, len(labels), lr, seed)
    loss_history, acc_history = model.fit(train_x, one_hot(train_y, len(labels)), epochs)
    train_pred = model.predict(train_x)
    test_pred = model.predict(test_x)
    train_acc = float((train_pred == train_y).mean())
    test_acc = float((test_pred == test_y).mean())
    cm = confusion_matrix(test_y, test_pred, len(labels))
    return RunResult(
        name=name,
        train_acc=train_acc,
        test_acc=test_acc,
        final_loss=loss_history[-1],
        epochs=epochs,
        hidden_size=hidden,
        learning_rate=lr,
        train_size=len(train_idx),
        test_size=len(test_idx),
        confusion=cm,
        labels=labels,
        loss_history=loss_history,
        acc_history=acc_history,
    )


def get_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def draw_training_curve(result: RunResult, path: Path) -> None:
    width, height = 900, 520
    margin = 70
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font = get_font(24)
    small = get_font(18)
    draw.text((width // 2 - 130, 20), f"{result.name} BP训练曲线", fill=(20, 20, 20), font=font)
    x0, y0 = margin, height - margin
    x1, y1 = width - margin, margin
    draw.line((x0, y0, x1, y0), fill=(0, 0, 0), width=2)
    draw.line((x0, y0, x0, y1), fill=(0, 0, 0), width=2)
    draw.text((x1 - 45, y0 + 16), "epoch", fill=(0, 0, 0), font=small)
    draw.text((12, y1 - 10), "loss", fill=(0, 0, 0), font=small)
    values = np.array(result.loss_history)
    max_v, min_v = float(values.max()), float(values.min())
    span = max(max_v - min_v, 1e-9)
    sample = np.linspace(0, len(values) - 1, min(300, len(values))).astype(int)
    points = []
    for i in sample:
        px = x0 + (x1 - x0) * i / (len(values) - 1)
        py = y0 - (y0 - y1) * (values[i] - min_v) / span
        points.append((px, py))
    draw.line(points, fill=(35, 102, 170), width=3)
    for ratio in np.linspace(0, 1, 5):
        ytick = y0 - (y0 - y1) * ratio
        val = min_v + span * ratio
        draw.line((x0 - 5, ytick, x0, ytick), fill=(0, 0, 0), width=1)
        draw.text((8, ytick - 10), f"{val:.3f}", fill=(0, 0, 0), font=small)
    img.save(path)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    set_run_font(run, "宋体", 12)


def set_run_font(run, font_name: str, size_pt: float) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)


def add_paragraph(doc: Document, text: str, first_line: bool = True) -> None:
    p = doc.add_paragraph()
    if first_line:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, "宋体", 12)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    set_run_font(run, "宋体", 14 if level == 1 else 12)


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "000000")
        borders.append(tag)
    tbl_pr.append(borders)


def build_report(results: list[RunResult], curve_paths: list[Path], report_path: Path) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.6)
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("实验二  BP算法实践实验报告")
    run.bold = True
    set_run_font(run, "宋体", 16)

    info = doc.add_table(rows=3, cols=4)
    set_table_borders(info)
    labels = [("课程名称", "机器学习"), ("实验名称", "BP算法实践"), ("实验工具", "Python、NumPy"), ("数据集", "鸢尾花数据集、西瓜二分类数据集"), ("学生姓名", "严浩睿"), ("学号", "20240395")]
    k = 0
    for r in range(3):
        for c in range(0, 4, 2):
            set_cell_text(info.cell(r, c), labels[k][0], True)
            set_cell_text(info.cell(r, c + 1), labels[k][1])
            k += 1

    add_heading(doc, "一、实验目的")
    add_paragraph(doc, "理解 BP 神经网络的基本结构、前向传播过程和误差反向传播思想，掌握使用 Python 从零实现多层前馈神经网络的方法。")
    add_paragraph(doc, "在鸢尾花数据集和工作区提供的二分类数据集上完成训练、预测与评价，观察学习率、隐含层规模和训练轮数对模型收敛效果的影响。")

    add_heading(doc, "二、实验原理")
    add_paragraph(doc, "BP 算法是一种基于梯度下降的神经网络训练方法。样本首先经过输入层到隐含层、隐含层到输出层的前向计算，得到类别概率；随后根据真实标签与预测结果之间的损失，按照链式法则逐层计算梯度，并更新权重和偏置。")
    add_paragraph(doc, "本实验采用单隐含层前馈神经网络。隐含层使用 Sigmoid 激活函数，输出层使用 Softmax 函数，多分类任务使用交叉熵损失。对二分类任务仍按两类 Softmax 处理，使两个实验共用同一套 BP 代码。")

    add_heading(doc, "三、实验环境与数据集")
    add_paragraph(doc, "实验程序使用 Python 编写，核心计算使用 NumPy 实现，没有调用现成的神经网络训练框架。为了提高收敛稳定性，训练前对连续特征进行标准化处理，并采用固定随机种子划分训练集和测试集。")
    table = doc.add_table(rows=1, cols=5)
    set_table_borders(table)
    headers = ["数据集", "样本数", "特征数", "类别数", "划分方式"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, True)
    rows = [
        ["鸢尾花", "150", "4", "3", "按类别分层，70% 训练、30% 测试"],
        ["二分类数据集", "17", "2", "2", "按类别分层，约 70% 训练、30% 测试"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)

    add_heading(doc, "四、程序设计与实现")
    add_paragraph(doc, "程序首先读取 CSV 数据，将类别标签转换为整数编码，再按类别进行分层抽样，保证训练集和测试集中各类别均有样本。随后使用训练集均值和标准差对特征做标准化，避免不同量纲造成梯度更新不稳定。")
    add_paragraph(doc, "网络参数包括输入层到隐含层权重 W1、隐含层偏置 b1、隐含层到输出层权重 W2 和输出层偏置 b2。每一轮训练依次完成前向传播、交叉熵损失计算、输出层梯度计算、隐含层梯度计算和参数更新。")
    add_paragraph(doc, "评价指标主要包括训练准确率、测试准确率、最终损失值和混淆矩阵。训练过程中记录每轮损失变化，用于判断模型是否收敛。")

    add_heading(doc, "五、实验结果与分析")
    summary = doc.add_table(rows=1, cols=7)
    set_table_borders(summary)
    for i, h in enumerate(["数据集", "隐含层节点", "学习率", "训练轮数", "训练准确率", "测试准确率", "最终损失"]):
        set_cell_text(summary.cell(0, i), h, True)
    for result in results:
        cells = summary.add_row().cells
        values = [
            result.name,
            str(result.hidden_size),
            f"{result.learning_rate:.3f}",
            str(result.epochs),
            f"{result.train_acc * 100:.2f}%",
            f"{result.test_acc * 100:.2f}%",
            f"{result.final_loss:.4f}",
        ]
        for i, value in enumerate(values):
            set_cell_text(cells[i], value)

    for result, curve in zip(results, curve_paths):
        add_heading(doc, f"{result.name}结果", level=2)
        add_paragraph(doc, f"{result.name}训练集样本数为 {result.train_size}，测试集样本数为 {result.test_size}。模型训练准确率为 {result.train_acc * 100:.2f}%，测试准确率为 {result.test_acc * 100:.2f}%，最终损失为 {result.final_loss:.4f}。")
        cm_table = doc.add_table(rows=len(result.labels) + 1, cols=len(result.labels) + 1)
        set_table_borders(cm_table)
        set_cell_text(cm_table.cell(0, 0), "真实\\预测", True)
        for i, label in enumerate(result.labels):
            set_cell_text(cm_table.cell(0, i + 1), label, True)
            set_cell_text(cm_table.cell(i + 1, 0), label, True)
        for i in range(len(result.labels)):
            for j in range(len(result.labels)):
                set_cell_text(cm_table.cell(i + 1, j + 1), str(int(result.confusion[i, j])))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(curve), width=Cm(13.5))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f"图 {results.index(result) + 1} {result.name}训练损失变化曲线")
        set_run_font(r, "宋体", 10.5)

    add_paragraph(doc, "从实验结果可以看出，BP 神经网络能够较好地学习鸢尾花数据集中花萼、花瓣特征与类别之间的非线性关系，测试集上保持较高准确率。二分类数据集样本量较小，模型对训练集划分较敏感，但损失曲线整体下降，说明反向传播更新方向有效。")
    add_paragraph(doc, "与线性分类器相比，BP 网络通过隐含层引入非线性变换，表达能力更强；但模型参数较多，在小样本数据上容易受初始权重和数据划分影响。因此实际应用中应结合交叉验证、正则化或更多样本来提高泛化能力。")

    add_heading(doc, "六、实验结论")
    add_paragraph(doc, "本实验完成了基于 Python 和 NumPy 的 BP 神经网络分类程序，实现了数据读取、标准化、前向传播、误差反向传播、参数更新和模型评价等完整流程。实验结果表明，BP 算法能够通过多轮迭代逐步降低损失，并在鸢尾花等典型分类任务上取得较好的分类效果。")
    add_paragraph(doc, "通过本次实验，我进一步理解了神经网络训练的核心并不只是调用模型接口，而是由损失函数、梯度计算、激活函数和参数更新共同构成的优化过程。后续若继续改进，可尝试不同隐含层节点数、学习率、批量训练策略和正则化方法，以比较其对模型泛化性能的影响。")

    add_heading(doc, "附录：核心程序说明")
    add_paragraph(doc, "完整程序文件为 bp_experiment_report.py。核心类 BPClassifier 封装了 Sigmoid、Softmax、前向传播、fit 训练和 predict 预测函数；run_dataset 函数负责完成单个数据集的训练和评价；build_report 函数负责生成实验报告文档。")

    doc.save(report_path)


def main() -> None:
    iris_x, iris_y, iris_labels = read_iris()
    wm_x, wm_y, wm_labels = read_watermelon()
    iris = run_dataset("鸢尾花数据集", iris_x, iris_y, iris_labels, hidden=10, lr=0.08, epochs=3500, test_ratio=0.30, seed=7)
    watermelon = run_dataset("二分类数据集", wm_x, wm_y, wm_labels, hidden=6, lr=0.12, epochs=2500, test_ratio=0.30, seed=2)
    results = [iris, watermelon]
    curve_paths = []
    for result in results:
        path = OUT_DIR / f"{result.name}_loss.png"
        draw_training_curve(result, path)
        curve_paths.append(path)
    report_path = ROOT / "实验2-BP算法实践-已完成.docx"
    build_report(results, curve_paths, report_path)
    for result in results:
        print(f"{result.name}: train_acc={result.train_acc:.4f}, test_acc={result.test_acc:.4f}, loss={result.final_loss:.4f}")
        print(result.confusion)
    print(report_path)


if __name__ == "__main__":
    main()
